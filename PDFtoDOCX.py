import fitz  # PyMuPDF
from docx import Document

# Define file paths
pdf_file = r'C:\path\to\your\cv.pdf'  # Replace with your PDF path
docx_file = r'C:\path\to\your\cv.docx'  # Replace with your desired DOCX path


# Initialize the PDF document
pdf_document = fitz.open(pdf_file)
doc = Document()

try:
    # Extract text from each page of the PDF
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)  # Load each page
        text = page.get_text("text")  # Extract text
        
        # Add extracted text to DOCX file
        doc.add_paragraph(text)
        doc.add_paragraph("\n")  # Add a newline after each page's content

    # Save the document as DOCX
    doc.save(docx_file)

    print(f"PDF text extracted and saved as DOCX: {docx_file}")

except Exception as e:
    print(f"An error occurred: {e}")

# Wait for input to see the error message # In case on error Programm shows you without closing
input("Press Enter to exit...")
